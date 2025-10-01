import json
from pathlib import Path
from typing import List, Dict, Tuple
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def extract_card_statuses_by_list(
    file_path: Path,
    target_lists: List[str]
) -> Dict[str, List[Tuple[str, str]]]:
    try:
        with file_path.open(encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"File not found: {file_path}")
        return {}
    except json.JSONDecodeError:
        print(f"Invalid JSON in file: {file_path}")
        return {}

    # Map list IDs to names
    list_id_to_name = {
        lst['id']: lst['name']
        for lst in data.get('lists', [])
        if lst.get('name') in target_lists
    }

    cards = data.get('cards', [])
    grouped_results: Dict[str, List[Tuple[str, str]]] = {}

    for card in cards:
        if card.get('closed', False):  # Skip archived cards
            continue
        list_id = card.get('idList')
        list_name = list_id_to_name.get(list_id)

        if not list_name:
            continue  # Skip cards in lists not in target_lists

        title = str(card.get('name', '')).strip()
        desc = str(card.get('desc', '')).strip()
        if not (title and desc):
            continue

        first_line = desc.splitlines()[0].strip()
        if not first_line:
            continue

        grouped_results.setdefault(list_name, []).append((title, first_line))

    return grouped_results

def export_to_word(grouped: Dict[str, List[Tuple[str, str]]], output_path: Path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for list_name in grouped:
        cards = grouped[list_name]
        if not cards:
            continue

        para = doc.add_paragraph()
        run = para.add_run(f"== {list_name} ==")
        run.bold = True

        for title, status in cards:
            para = doc.add_paragraph(style='Normal')
            run = para.add_run(f"• {title}")
            run.bold = True

            para = doc.add_paragraph(style='Normal')
            para.paragraph_format.left_indent = Pt(18)
            run = para.add_run(f"o Status: ")
            run.bold = True
            para.add_run(status)

        doc.add_paragraph("")  # Add space between lists

    doc.save(output_path)

def main():
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
    else:
        path = Path("example_data/trello_export.json")  # Replace with your actual JSON file

    target_lists = [
        "List1", "List2", "List3",
        "List4"] # Replace with Trello lists that you want to extract

    grouped = extract_card_statuses_by_list(path, target_lists)

    output_path = Path("trello_status_report.docx")
    export_to_word(grouped, output_path)
    print(f"Exported Word report to {output_path}")

if __name__ == "__main__":
    main()

# This code extracts card statuses from a Trello JSON file and groups them by specified lists.
# It prints the title and first line of the description for each card in the specified lists.
# It also exports the formatted results into a Word document with bolded headers and structured indentation.
 
